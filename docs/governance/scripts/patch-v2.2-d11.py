#!/usr/bin/env python3
"""Idempotent D11 § Period States patch for Constitution v2.2 DOCX + merge report."""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

EXPECTED = (
    "The official period states are OPEN, CLOSING, and CLOSED. "
    "The state Archived is not a period registry state; historical snapshots and reports "
    "use SUPERSEDED versioning (§6.11, §6.17)."
)
FORBIDDEN = "Open, Closing, Closed, Archived"
HEADING = "6.4 Period States"


def _patch_docx(path: Path) -> bool:
    from docx import Document

    doc = Document(str(path))
    changed = False
    heading_idx = None
    for i, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if text == HEADING:
            heading_idx = i
        if FORBIDDEN in text or text == "Open, Closing, Closed, Archived.":
            para.text = EXPECTED
            changed = True
        if text == EXPECTED:
            # already correct
            pass

    if heading_idx is not None:
        # Ensure the paragraph immediately after the heading is the approved wording.
        body_idx = heading_idx + 1
        while body_idx < len(doc.paragraphs) and not (doc.paragraphs[body_idx].text or "").strip():
            body_idx += 1
        if body_idx < len(doc.paragraphs):
            current = (doc.paragraphs[body_idx].text or "").strip()
            if current != EXPECTED:
                doc.paragraphs[body_idx].text = EXPECTED
                changed = True
        else:
            doc.add_paragraph(EXPECTED)
            changed = True

    if changed:
        doc.save(str(path))
    return changed


def _patch_merge_report(path: Path) -> bool:
    text = path.read_text(encoding="utf-8")
    original = text
    if FORBIDDEN in text:
        text = text.replace(FORBIDDEN + ".", EXPECTED)
        text = text.replace(FORBIDDEN, EXPECTED)
    # Ensure D11 note exists near top decisions if report uses that format.
    if EXPECTED not in text:
        raise SystemExit(f"{path.name}: cannot locate/insert expected D11 wording automatically")
    # Normalize §6.4 Period States section body when present.
    text = re.sub(
        r"(6\.4 Period States\s*\n)(.*?)(\n6\.5 |\n## |\Z)",
        lambda m: f"{m.group(1)}{EXPECTED}\n{m.group(3)}",
        text,
        count=1,
        flags=re.S,
    )
    if text != original:
        path.write_text(text, encoding="utf-8", newline="\n")
        return True
    return False


def main() -> int:
    root = Path(__file__).resolve().parents[1]
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", type=Path, default=root / "DX_OSE_CONSTITUTION_v2.2.docx")
    parser.add_argument(
        "--merge-report",
        type=Path,
        default=root / "DX_OSE_CONSTITUTION_v2.1_MERGE_REPORT.md",
    )
    args = parser.parse_args()

    results = {
        "docx_changed": _patch_docx(args.docx),
        "merge_report_changed": _patch_merge_report(args.merge_report),
    }
    print(results)
    return 0


if __name__ == "__main__":
    sys.exit(main())
