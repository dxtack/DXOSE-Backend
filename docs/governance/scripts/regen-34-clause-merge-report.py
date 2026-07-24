"""Regenerate DX_OSE_34_Clause_Merge_Report_v1.docx after Path A renumbering."""
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parents[1] / "DX_OSE_34_Clause_Merge_Report_v1.docx"


def add_p(doc, text, bold=False, size=9, name="Aptos"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    return p


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.5)

    add_p(doc, "DX OSE — 34-Clause Merge Report v1", bold=True, size=14)
    add_p(
        doc,
        "Separate merge evidence for DX_OSE_CONSTITUTION_v2.1_MERGED / v2.2 "
        "(Path A Chapter 6 renumber applied).",
    )
    add_p(
        doc,
        "Classification rules: موجود مسبقًا = already governed; اتضاف = inserted; "
        "تعارض اتحل = conflict resolved using approved recommendation.",
    )
    doc.add_paragraph()
    add_p(doc, "# | Merge status | Decision / evidence", bold=True, size=8)

    rows = [
        (1, "اتضاف", "§4.5.1 ADMIN retirement."),
        (
            2,
            "تعارض اتحل",
            "§4.5.2 Full ACC means full operational grants excluding period-close/reopen, "
            "platform-only, currency configuration and Master Data mutation.",
        ),
        (
            3,
            "موجود مسبقًا",
            "ORG_MANAGER is recognized in approved access/period governance; "
            "no contradictory role removal exists.",
        ),
        (4, "اتضاف", "§11.9 Display-currency mutation reserved to SUPER_ADMIN."),
        (5, "اتضاف", "§4.5.4 Master Data mutation allowlist."),
        (6, "اتضاف", "§4.5.5 GENERAL_MANAGER exclusions."),
        (
            7,
            "تعارض اتحل",
            "§6.10 PC-IC warning alone is nonblocking; an incomplete count integrity "
            "condition remains a blocker. (Path A: former §6.16)",
        ),
        (8, "اتضاف", "§3.9 Unified mandatory Send Back target list."),
        (
            9,
            "موجود مسبقًا",
            "Posted business outcome and status consistency governed by §§2.2–2.3; "
            "merged text adds explicit implementation-facing integrity wording.",
        ),
        (
            10,
            "تعارض اتحل",
            "§6.15 UTC internal persistence retained; tenant timezone defines business EOD. "
            "(Path A: former §6.21)",
        ),
        (
            11,
            "موجود مسبقًا",
            "Approved Chapter 10 integer-quantity rule; the merged DOCX carries it as §10.5 "
            "because v2.0 source had not consolidated it.",
        ),
        (
            12,
            "موجود مسبقًا",
            "No page-level horizontal scrolling except internal grids; broader registry rule "
            "exists in consolidated Chapter 30 (reserved / not merged in official DOCX).",
        ),
        (13, "اتضاف", "§23.8 Location-constrained item lookup."),
        (14, "اتضاف", "§32.2 Process Return restricted to the three governed roles."),
        (
            15,
            "اتضاف",
            "§32.3 Force Close remains a protected service capability without enabled "
            "interactive action.",
        ),
        (16, "اتضاف", "§32.4 Damaged/Lost dispositions follow BRK/LST governance."),
        (17, "اتضاف", "§32.5 Immediate Good-return posting."),
        (18, "اتضاف", "§32.6 Three-way quantity split and outstanding limit."),
        (19, "اتضاف", "§32.7 Optional Damaged-only photos."),
        (20, "اتضاف", "§32.8 Dynamic Process Return / Returned behavior."),
        (21, "اتضاف", "§15.7 Original-note preservation."),
        (
            22,
            "موجود مسبقًا",
            "Every workflow action must create an auditable chronological event under "
            "§§22.2–22.3.",
        ),
        (23, "اتضاف", "§32.9 Complete inter-hotel lifecycle."),
        (
            24,
            "اتضاف",
            "§32.10 Ordinary Process Return prohibited for inter-hotel passes.",
        ),
        (
            25,
            "موجود مسبقًا",
            "Approved Chapter 10 WAC rule; merged DOCX carries it as §10.4.",
        ),
        (
            26,
            "موجود مسبقًا",
            "Posting Date determines the assigned period under Chapter 6.",
        ),
        (
            27,
            "موجود مسبقًا",
            "Approved Period Registry governance prohibits implicit open; merged DOCX "
            "carries it as §6.11. (Path A: former §6.17)",
        ),
        (
            28,
            "موجود مسبقًا",
            "Approved close lifecycle creates the snapshot only on successful close; "
            "merged DOCX carries it as §6.12. (Path A: former §6.18)",
        ),
        (
            29,
            "اتضاف",
            "§6.13 Closing snapshot as next opening reference. (Path A: former §6.19)",
        ),
        (
            30,
            "اتضاف",
            "§6.14 Close neither validates nor creates the following period. "
            "(Path A: former §6.20)",
        ),
        (
            31,
            "تعارض اتحل وظيفيًا",
            "§6.16 Granular Close Authority — PERIOD_CLOSE_EXECUTE canonical; "
            "PERIOD_CLOSE_MANAGE temporary compatibility. "
            "(Path A: former combined §6.22 close half)",
        ),
        (
            32,
            "تعارض اتحل وظيفيًا",
            "§6.17 Granular Reopen Authority — PERIOD_REOPEN_EXECUTE plus reason; "
            "MANAGE is not permanent authority. "
            "(Path A: former combined §6.22 reopen half; duplicate §6.22 numbering removed)",
        ),
        (
            33,
            "اتضاف",
            "§6.18 Governed seven-step monthly sequence. (Path A: former §6.23)",
        ),
        (
            34,
            "موجود مسبقًا",
            "Accepted inventory-truth decision confirms live Current Stock Balance; "
            "merged DOCX carries it as §33.1.",
        ),
    ]
    for n, status, evidence in rows:
        add_p(doc, f"{n}  {status} — {evidence}", size=8)

    doc.add_paragraph()
    add_p(doc, "Numbering review note (P2 #32 — Path A applied)", bold=True, size=11)
    notes = [
        "Path A renumbered merge-inserted Chapter 6 clauses contiguously after §6.9: "
        "former §6.16–§6.23 → current §6.10–§6.18 (with former combined §6.22 split into "
        "§6.16 close + §6.17 reopen). Numbering only; normative decisions unchanged.",
        "Deferred — D11: §6.4 still cites SUPERSEDED versioning as (§6.11, §6.17). Those "
        "citations intend consolidated Snapshot Versioning / Report Versioning content "
        "that is not present in the currently ratified DOCX. They must not be read as the "
        "Path A ratified §6.11 (Explicit Period Opening) or §6.17 (Granular Reopen "
        "Authority). Citation rewrite and any Snapshot/Report Versioning merge are "
        "deferred to a later review.",
        "Reserved / not merged — Chapters 30 (Unified Visual Language) and 31 "
        "(Repository Root Governance) remain outside the official ratified constitution "
        "DOCX. Chapters 32 and 33 retain their current numbers.",
    ]
    for note in notes:
        add_p(doc, "• " + note, size=9)

    doc.save(OUT)
    print(f"wrote {OUT} ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
